import streamlit as st
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import streamlit.components.v1 as components
import json
import gspread

st.set_page_config(page_title="El Silencio de Gabriela - Plan Lector", page_icon="📖", layout="centered")

# --- CONEXIÓN A GOOGLE SHEETS Y CANDADO DE 3 CERROJOS ---
@st.cache_resource
def get_google_sheet():
    creds_json = st.secrets["google_credentials"]
    creds_dict = json.loads(creds_json)
    if isinstance(creds_dict, list): creds_dict = creds_dict[0]
    client = gspread.service_account_from_dict(creds_dict)
    return client.open("Registro_Plan_Lector").sheet1

def ya_participo(nombre, inst, obra):
    try:
        sheet = get_google_sheet()
        data = sheet.get_all_values()
        if not data: return False
        for row in data[1:]:
            if len(row) >= 9:
                if (row[1].strip().lower() == nombre.strip().lower() and 
                    row[2].strip().lower() == inst.strip().lower() and 
                    row[8].strip().lower() == obra.strip().lower()):
                    return True
        return False
    except Exception as e:
        return False 

def registrar_y_bloquear(nombre, inst, nivel, grado, seccion, area, prof, obra):
    st.session_state.ficha_completada_cat4 = True
    try:
        sheet = get_google_sheet()
        fecha_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        sheet.append_row([fecha_str, nombre.upper(), inst.upper(), nivel, grado, seccion, area.upper(), prof, obra.upper(), "COMPLETADO"])
    except Exception as e:
        pass

# --- FUNCIONES DE ESTADO Y VALIDACIÓN ---
if 'ficha_completada_cat4' not in st.session_state: st.session_state.ficha_completada_cat4 = False
if 'intento_descarga_cat4' not in st.session_state: st.session_state.intento_descarga_cat4 = False

if st.session_state.ficha_completada_cat4:
    st.success("✅ **¡PRUEBA ENTREGADA CON ÉXITO!**")
    st.info("Has generado tu documento correctamente y tu participación ya está registrada de forma segura en el colegio.")
    st.write("Por favor, envía el archivo Word que se descargó a tu profesor(a). ¡Buen trabajo!")
    st.stop() # Esto es lo que detiene la página de forma invisible

def contiene_spam(texto):
    for palabra in texto.split():
        if len(palabra) > 20: return True
    return False

# Mínimo de palabras: Nivel avanzado (Micro-ensayos)
def obtener_minimos(grado):
    if grado == "4to": return {"q3": 20, "q4": 25, "q5": 30}
    elif grado == "5to": return {"q3": 25, "q4": 30, "q5": 40}
    return {"q3": 20, "q4": 20, "q5": 20}

# --- 1. LA BÓVEDA DE SEGURIDAD ---
nombre_de_la_obra = "El Silencio de Gabriela"

if 'autenticado_cat4' not in st.session_state: st.session_state.autenticado_cat4 = False

if not st.session_state.autenticado_cat4:
    st.title("📖 Ecosistema Literario Transmedia")
    st.subheader(f"Obra: {nombre_de_la_obra}")
    st.write("Autor: Eduardo Florez Montero")
    st.markdown("---")
    st.info("🔒 **Acceso Exclusivo:** Escribe la Clave de Aula que te indicó tu profesor.")
    
    clave = st.text_input("🔑 Clave de Acceso *:", type="password", autocomplete="new-password")
    
    if st.button("Desbloquear Evaluación", use_container_width=True):
        claves_validas = ["GABRIELA-2026", "SANJUAN-4A", "SANJUAN-5B"]
        if clave in claves_validas: 
            st.session_state.autenticado_cat4 = True
            st.rerun()
        else:
            st.error("❌ Clave incorrecta. Consulta con tu profesor.")
    st.stop()

# --- 2. SISTEMA DE TIEMPO INTELIGENTE ---
if 'ficha_iniciada_cat4' not in st.session_state: st.session_state.ficha_iniciada_cat4 = False
if 'extra_time_used_cat4' not in st.session_state: st.session_state.extra_time_used_cat4 = False

if not st.session_state.ficha_iniciada_cat4:
    st.success("✅ ¡Clave aceptada! Bienvenido al ecosistema.")
    st.warning("⚠️ **ATENCIÓN: TIENES UN SOLO INTENTO**\nEl sistema registrará tu nombre y colegio. No podrás volver a dar la evaluación de este libro.")
    
    try:
        st.image("portada_gabriela.jpg", width=250)
    except:
        st.caption("(Aquí aparecerá la portada de tu libro)")
        
    st.info("👋 Tienes 20 minutos para resolver tu evaluación de nivel avanzado.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO, INICIAR EVALUACIÓN", use_container_width=True, type="primary"):
            st.session_state.ficha_iniciada_cat4 = True
            st.session_state.inicio_tiempo_cat4 = time.time()
            st.session_state.minutos_asignados_cat4 = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_cat4
segundos_restantes = (st.session_state.minutos_asignados_cat4 * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_preguntas = tiempo_agotado 

with st.sidebar:
    st.markdown("### ⏱️ Tiempo Restante")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <div id="aviso_30s" style="color: #e74c3c; font-weight: bold; font-size: 14px; text-align: center; margin-top: 10px; display: none;">
            ⚠️ ¡Aviso! Te quedan 30 segundos.
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var aviso = document.getElementById('aviso_30s');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                    aviso.style.display = "block";
                    aviso.innerHTML = "⚠️ Tiempo agotado.";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; 
                    if (tiempo <= 60) display.style.color = "#e74c3c"; 
                    if (tiempo <= 30) aviso.style.display = "block";
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=130)
        
        if not st.session_state.extra_time_used_cat4:
            if st.button("➕ Dar 4 min extra", use_container_width=True):
                st.session_state.minutos_asignados_cat4 += 4
                st.session_state.extra_time_used_cat4 = True
                st.rerun()
        else:
            st.caption("✔️ Tiempo extra utilizado.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Aún puedes descargar tu avance.")

# --- 3. CONTENIDO LITERARIO Y VALIDACIONES ---
st.markdown(f"### 📖 {nombre_de_la_obra.upper()}")
st.markdown("**Autor:** Eduardo Florez Montero | **Categoría:** Jóvenes Adultos (4to y 5to Sec.)")
st.markdown("---")

css_error = "<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Este campo es obligatorio</div>"

institucion = st.text_input("Institución Educativa (Nombre completo) *:", disabled=False)
if st.session_state.intento_descarga_cat4 and not institucion.strip(): st.markdown(css_error, unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1: 
    nombre = st.text_input("Estudiante (Apellidos y Nombres) *:", disabled=False)
    if st.session_state.intento_descarga_cat4 and not nombre.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col2: 
    nivel = st.selectbox("Nivel *:", ["", "Secundaria"], disabled=False)
    if st.session_state.intento_descarga_cat4 and nivel == "": st.markdown(css_error, unsafe_allow_html=True)
with col3: 
    # SOLO 4TO Y 5TO
    grado = st.selectbox("Grado *:", ["", "4to", "5to"], disabled=False)
    if st.session_state.intento_descarga_cat4 and grado == "": st.markdown(css_error, unsafe_allow_html=True)

col4, col5, col6 = st.columns(3)
with col4: 
    seccion = st.selectbox("Sección *:", ["", "A", "B", "C", "D", "E", "F", "G", "H", "Única"], disabled=False)
    if st.session_state.intento_descarga_cat4 and seccion == "": st.markdown(css_error, unsafe_allow_html=True)
with col5: 
    area_curso = st.text_input("Área o Curso (Ej: Comunicación) *:", disabled=False)
    if st.session_state.intento_descarga_cat4 and not area_curso.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col6: 
    fecha_input = st.date_input("Fecha *:", disabled=False)

profesor = st.text_input("Nombre de tu Profesor(a) (Opcional):", disabled=False)

st.markdown("---")
st.subheader("🎬 Antes de empezar: Resumen Visual")
st.video("https://www.youtube.com/watch?v=123456789") 
st.markdown("---")

# --- LÓGICA DE DIFERENCIACIÓN COGNITIVA AVANZADA ---
if grado == "4to":
    p1 = "1) Sintetiza el argumento principal de la obra identificando el suceso que detona el problema: *"
    p2 = "2) ¿Cómo influye el contexto sociocultural de la historia en las decisiones del personaje principal? *"
    p3 = "3) Analiza el dilema moral que enfrenta el protagonista. ¿Existen 'buenos' y 'malos' absolutos en esta obra? Argumenta. *"
    p4 = "4) Explica cómo se resuelve el conflicto principal y qué consecuencias trae para el entorno social de los personajes. *"
    p5 = "5) Redacta un breve ensayo crítico: ¿De qué manera la historia de Gabriela refleja problemáticas vigentes en la sociedad peruana actual? *"
    
    c1 = "NIVEL 1: Sintetiza el argumento y describe la influencia del contexto sociocultural en la trama."
    c2 = "NIVEL 2: Analiza matices morales y desglosa las consecuencias sistémicas del conflicto."
    c3 = "NIVEL 3: Redacta una reflexión crítica estructurada vinculando la literatura con la realidad nacional."

elif grado == "5to":
    p1 = "1) Describe la evolución psicológica del personaje principal desde el inicio hasta el desenlace de la obra: *"
    p2 = "2) Identifica y explica los símbolos o metáforas sociales que el autor utiliza para construir la atmósfera de la historia: *"
    p3 = "3) Desde una perspectiva filosófica o ética, ¿cómo juzgas la resolución del conflicto? Fundamenta tu postura. *"
    p4 = "4) Si fueras el autor (Eduardo Florez), ¿cuál dirías que fue tu intencionalidad principal al escribir esta obra? *"
    p5 = "5) Ensayo de síntesis: Propón una solución estructural, desde el rol ciudadano, al problema social que denuncia la historia de Gabriela. *"
    
    c1 = "NIVEL 1: Describe la evolución psicológica de los personajes e identifica elementos simbólicos en la narrativa."
    c2 = "NIVEL 2: Juzga dilemas éticos con fundamentación filosófica e infiere la intencionalidad del autor."
    c3 = "NIVEL 3: Construye un ensayo propositivo, formulando soluciones ciudadanas a problemáticas estructurales."

minimos = obtener_minimos(grado if grado != "" else "4to")

if grado != "":
    st.info(f"🧠 Evaluación Avanzada (Pensamiento Formal) adaptada para {grado} de Secundaria.")
    
    st.markdown("### 📋 Rúbrica de Evaluación Continua")
    st.success(f"**NIVEL 1 (Análisis Literario):** {c1}\n\n**NIVEL 2 (Inferencia Compleja):** {c2}\n\n**NIVEL 3 (Crítica y Propuesta):** {c3}")
    st.markdown("---")
    
    st.subheader("NIVEL 1: ANÁLISIS E INTERPRETACIÓN LITERARIA")
    q1 = st.text_area(p1, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga_cat4 and not tiempo_agotado and not q1.strip(): st.markdown(css_error, unsafe_allow_html=True)

    q2 = st.text_area(p2, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga_cat4 and not tiempo_agotado and not q2.strip(): st.markdown(css_error, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("NIVEL 2: INFERENCIA Y EVALUACIÓN ÉTICA")
    q3 = st.text_area(p3, disabled=bloquear_preguntas)
    palabras_q3 = len(q3.split())
    if st.session_state.intento_descarga_cat4 and not tiempo_agotado:
        if not q3.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q3 < minimos["q3"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Mínimo {minimos['q3']} palabras (tienes {palabras_q3}).</div>", unsafe_allow_html=True)

    q4 = st.text_area(p4, disabled=bloquear_preguntas)
    palabras_q4 = len(q4.split())
    if st.session_state.intento_descarga_cat4 and not tiempo_agotado:
        if not q4.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q4 < minimos["q4"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Mínimo {minimos['q4']} palabras (tienes {palabras_q4}).</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("NIVEL 3: MICRO-ENSAYO Y REFLEXIÓN CRÍTICA")
    q5 = st.text_area(p5, disabled=bloquear_preguntas, height=150)
    palabras_q5 = len(q5.split())
    if st.session_state.intento_descarga_cat4 and not tiempo_agotado:
        if not q5.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q5 < minimos["q5"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Mínimo {minimos['q5']} palabras (tienes {palabras_q5}).</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.error("🚨 **ÚLTIMO PASO:** Al generar el Word, tu nombre quedará registrado en el sistema central de control.")

    # --- LÓGICA DE VALIDACIÓN ---
    faltantes_personales = not institucion.strip() or not nombre.strip() or nivel == "" or grado == "" or seccion == "" or not area_curso.strip()
    preguntas_vacias = not q1.strip() or not q2.strip() or not q3.strip() or not q4.strip() or not q5.strip()
    faltan_palabras = palabras_q3 < minimos["q3"] or palabras_q4 < minimos["q4"] or palabras_q5 < minimos["q5"]
    hay_spam = contiene_spam(q1) or contiene_spam(q2) or contiene_spam(q3) or contiene_spam(q4) or contiene_spam(q5)

    hay_errores = False
    if faltantes_personales: hay_errores = True
    if not tiempo_agotado and (preguntas_vacias or faltan_palabras or hay_spam): hay_errores = True

    if not st.session_state.intento_descarga_cat4 or hay_errores:
        if st.button("📥 Generar y Descargar mi Ensayo", type="primary", use_container_width=True):
            st.session_state.intento_descarga_cat4 = True
            st.rerun()
        
        if st.session_state.intento_descarga_cat4 and hay_errores:
            if faltantes_personales: st.error("⚠️ Faltan datos obligatorios. Revisa las alertas en rojo arriba.")
            elif not tiempo_agotado:
                if preguntas_vacias: st.warning("⚠️ Te falta responder algunas preguntas analíticas.")
                if faltan_palabras: st.warning("⚠️ Algunas respuestas no cumplen con la extensión mínima requerida para tu grado.")
                if hay_spam: st.error("⚠️ Escribe con normalidad, sin caracteres repetidos.")

    else:
        with st.spinner("Validando usuario en la Base de Datos..."):
            duplicado = ya_participo(nombre, institucion, nombre_de_la_obra)
            
        if duplicado:
            st.error(f"🛑 **BLOQUEO:** El estudiante '{nombre.upper()}' de la institución '{institucion.upper()}' ya rindió la evaluación de '{nombre_de_la_obra}'.")
        else:
            if tiempo_agotado:
                st.info("⏱️ El tiempo se agotó. Descargando avance parcial...")
            else:
                st.success("✨ ¡Todo correcto! Haz clic abajo para descargar tu evaluación crítica.")
                
            doc = Document()
            doc.add_heading('EVALUACIÓN DEL PLAN LECTOR (AVANZADO)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'I.E.: {institucion.upper()}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'Obra: {nombre_de_la_obra} | Autor: Eduardo Florez Montero').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            fecha_str = fecha_input.strftime("%d/%m/%Y")
            texto_area = area_curso.strip()
            texto_prof = profesor.strip() if profesor.strip() else "No especificado"

            doc.add_paragraph(f'Estudiante: {nombre}\nNivel: {nivel} | Grado: {grado} | Sección: {seccion} | Fecha: {fecha_str}')
            doc.add_paragraph(f'Área/Curso: {texto_area} | Docente a cargo: {texto_prof}')
            
            # --- FUNCIÓN AYUDANTE PARA JUSTIFICAR TEXTO ---
            def agregar_bloque_justificado(pregunta, respuesta):
                p_pregunta = doc.add_paragraph()
                p_pregunta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_pregunta.add_run(pregunta).bold = True
                
                p_respuesta = doc.add_paragraph()
                p_respuesta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                texto_respuesta = respuesta if respuesta.strip() else "[No respondió]"
                p_respuesta.add_run(f'Respuesta: {texto_respuesta}\n')

            # --- APLICANDO EL JUSTIFICADO A LOS NIVELES ---
            doc.add_heading('NIVEL 1 (Análisis Literario)', level=2)
            agregar_bloque_justificado(p1, q1)
            agregar_bloque_justificado(p2, q2)
            
            doc.add_heading('NIVEL 2 (Inferencia Compleja)', level=2)
            agregar_bloque_justificado(p3, q3)
            agregar_bloque_justificado(p4, q4)
            
            doc.add_heading('NIVEL 3 (Micro-Ensayo Crítico)', level=2)
            agregar_bloque_justificado(p5, q5)
            
            doc.add_page_break()
            doc.add_heading(f'Lista de Cotejo ({grado} Secundaria)', level=2)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = 'CRITERIOS DE EVALUACIÓN', 'INICIO', 'PROCESO', 'LOGRADO', 'DESTACADO'
            
            criterios = [c1, c2, c3] 
            
            for crit in criterios:
                row_cells = table.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text, row_cells[4].text = crit, "[   ]", "[   ]", "[   ]", "[   ]"
                
            doc.add_paragraph('\nObservaciones / Feedback al estudiante:\n________________________________________________')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 Confirmar Registro y Descargar Evidencia", 
                data=bio.getvalue(), 
                file_name=f"Evaluacion_{grado}_{seccion}_{nombre.replace(' ', '_')}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                on_click=registrar_y_bloquear,
                args=(nombre, institucion, nivel, grado, seccion, area_curso, profesor, nombre_de_la_obra),
                type="primary",
                use_container_width=True
            )
else:
    st.info("👆 Por favor, selecciona tu **Nivel** y **Grado** en la parte superior para cargar las preguntas de tu evaluación.")
