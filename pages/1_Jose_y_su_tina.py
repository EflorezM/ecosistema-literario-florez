import streamlit as st
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import streamlit.components.v1 as components
import json
import gspread

st.set_page_config(page_title="José y su tina océano - Plan Lector", page_icon="🧸", layout="centered")

# --- CONEXIÓN A GOOGLE SHEETS Y CANDADO ---
@st.cache_resource
def get_google_sheet():
    creds_json = st.secrets["google_credentials"]
    creds_dict = json.loads(creds_json)
    if isinstance(creds_dict, list): creds_dict = creds_dict[0]
    client = gspread.service_account_from_dict(creds_dict)
    return client.open("Registro_Plan_Lector").sheet1

def ya_participo(nombre, inst, obra):
    try:
        sheet = get_google_sheet()
        data = sheet.get_all_values()
        if not data: return False
        for row in data[1:]:
            if len(row) >= 9:
                if (row[1].strip().lower() == nombre.strip().lower() and 
                    row[2].strip().lower() == inst.strip().lower() and 
                    row[8].strip().lower() == obra.strip().lower()):
                    return True
        return False
    except Exception as e:
        return False 

def registrar_y_bloquear(nombre, inst, nivel, grado, seccion, area, prof, obra):
    st.session_state.ficha_completada_cat1 = True
    try:
        sheet = get_google_sheet()
        fecha_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        sheet.append_row([fecha_str, nombre.upper(), inst.upper(), nivel, grado, seccion, area.upper(), prof, obra.upper(), "COMPLETADO"])
    except Exception as e:
        pass

# --- FUNCIONES DE ESTADO Y VALIDACIÓN ---
if 'ficha_completada_cat1' not in st.session_state: st.session_state.ficha_completada_cat1 = False
if 'intento_descarga_cat1' not in st.session_state: st.session_state.intento_descarga_cat1 = False

if st.session_state.ficha_completada_cat1:
    st.success("✅ **¡TAREA ENTREGADA CON ÉXITO!**")
    st.info("¡Felicidades! Has terminado tu evaluación y tu trabajo ya está guardado seguro.")
    st.write("Por favor, dile a tus papis que envíen el archivo Word a tu profesor(a). ¡Eres un campeón/a!")
    st.stop()

def contiene_spam(texto):
    for palabra in texto.split():
        if len(palabra) > 20: return True
    return False

# Mínimo de palabras: Adaptado a primeros lectores (Piaget)
def obtener_minimos(grado):
    if grado == "1ro": return {"q3": 3, "q4": 3, "q5": 3}
    elif grado == "2do": return {"q3": 5, "q4": 5, "q5": 5}
    elif grado == "3ro": return {"q3": 8, "q4": 8, "q5": 8}
    return {"q3": 3, "q4": 3, "q5": 3}

# --- 1. LA BÓVEDA DE SEGURIDAD ---
nombre_de_la_obra = "José y su tina océano"

if 'autenticado_cat1' not in st.session_state: st.session_state.autenticado_cat1 = False

if not st.session_state.autenticado_cat1:
    st.title("🧸 Club de Pequeños Lectores")
    st.subheader(f"📖 Cuento: {nombre_de_la_obra}")
    st.write("Autor: Eduardo Florez Montero")
    st.markdown("---")
    st.info("🤫 **Palabra Mágica:** Escribe la clave que te dio tu profesor(a).")
    
    clave = st.text_input("🔑 Clave Mágica *:", type="password", autocomplete="new-password")
    
    if st.button("¡Entrar al cuento!", use_container_width=True):
        claves_validas = ["TINA-2026", "PRIMARIA-1A", "PRIMARIA-2B"]
        if clave in claves_validas: 
            st.session_state.autenticado_cat1 = True
            st.rerun()
        else:
            st.error("❌ Clave incorrecta. Pídele ayuda a tu profesor(a) o a tus papis.")
    st.stop()

# --- 2. SISTEMA DE TIEMPO INTELIGENTE ---
if 'ficha_iniciada_cat1' not in st.session_state: st.session_state.ficha_iniciada_cat1 = False
if 'extra_time_used_cat1' not in st.session_state: st.session_state.extra_time_used_cat1 = False

if not st.session_state.ficha_iniciada_cat1:
    st.success("✅ ¡Clave correcta! Bienvenido a esta aventura.")
    st.warning("⚠️ **ATENCIÓN:** Solo puedes hacer esta prueba una vez. ¡Hazlo con calma!")
    
    try:
        st.image("portada_jose.jpg", width=250)
    except:
        st.caption("(Aquí aparecerá la portada de tu cuento)")
        
    st.info("👋 Tienes 20 minutos para responder. Puedes pedirle a tus papis que te ayuden a escribir, ¡pero las ideas deben ser tuyas!")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ¡ESTOY LISTO, EMPEZAR!", use_container_width=True, type="primary"):
            st.session_state.ficha_iniciada_cat1 = True
            st.session_state.inicio_tiempo_cat1 = time.time()
            st.session_state.minutos_asignados_cat1 = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_cat1
segundos_restantes = (st.session_state.minutos_asignados_cat1 * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_preguntas = tiempo_agotado 

with st.sidebar:
    st.markdown("### ⏱️ Reloj Mágico")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <div id="aviso_30s" style="color: #e74c3c; font-weight: bold; font-size: 14px; text-align: center; margin-top: 10px; display: none;">
            ⚠️ ¡Aviso! Queda muy poquito tiempo.
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var aviso = document.getElementById('aviso_30s');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                    aviso.style.display = "block";
                    aviso.innerHTML = "⚠️ El tiempo terminó.";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; 
                    if (tiempo <= 60) display.style.color = "#e74c3c"; 
                    if (tiempo <= 30) aviso.style.display = "block";
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=130)
        
        if not st.session_state.extra_time_used_cat1:
            if st.button("➕ Pedir 4 minutos más", use_container_width=True):
                st.session_state.minutos_asignados_cat1 += 4
                st.session_state.extra_time_used_cat1 = True
                st.rerun()
        else:
            st.caption("✔️ Ya pediste tiempo extra.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO TERMINADO")
        st.write("Aún puedes descargar tu avance abajo.")

# --- 3. CONTENIDO LITERARIO Y VALIDACIONES ---
st.markdown(f"### 🧸 {nombre_de_la_obra.upper()}")
st.markdown("**Autor:** Eduardo Florez Montero | **Categoría:** Primeros Lectores (1ro a 3ro Primaria)")
st.markdown("---")

css_error = "<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ No olvides llenar esto</div>"

st.info("🏫 **Tus Datos (Pide ayuda si la necesitas)**")

institucion = st.text_input("Nombre de tu Colegio *:", disabled=False)
if st.session_state.intento_descarga_cat1 and not institucion.strip(): st.markdown(css_error, unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1: 
    nombre = st.text_input("Tus Nombres y Apellidos *:", disabled=False)
    if st.session_state.intento_descarga_cat1 and not nombre.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col2: 
    nivel = st.selectbox("Nivel *:", ["", "Primaria"], disabled=False)
    if st.session_state.intento_descarga_cat1 and nivel == "": st.markdown(css_error, unsafe_allow_html=True)
with col3: 
    # SOLO 1RO, 2DO Y 3RO
    grado = st.selectbox("Tu Grado *:", ["", "1ro", "2do", "3ro"], disabled=False)
    if st.session_state.intento_descarga_cat1 and grado == "": st.markdown(css_error, unsafe_allow_html=True)

col4, col5, col6 = st.columns(3)
with col4: 
    seccion = st.selectbox("Sección *:", ["", "A", "B", "C", "D", "E", "F", "G", "H", "Única"], disabled=False)
    if st.session_state.intento_descarga_cat1 and seccion == "": st.markdown(css_error, unsafe_allow_html=True)
with col5: 
    area_curso = st.text_input("Curso (Ej: Comunicación) *:", disabled=False)
    if st.session_state.intento_descarga_cat1 and not area_curso.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col6: 
    fecha_input = st.date_input("Fecha *:", disabled=False)

profesor = st.text_input("Nombre de tu Profesor(a) (Opcional):", disabled=False)

st.markdown("---")
st.subheader("🎬 Sorpresa antes de empezar")
st.video("https://www.youtube.com/watch?v=123456789") 
st.markdown("---")

# --- LÓGICA DE DIFERENCIACIÓN COGNITIVA (PRIMARIA) ---
if grado == "1ro":
    p1 = "1) ¿Cómo se llama el personaje principal del cuento? *"
    p2 = "2) ¿En qué lugar de la casa jugaba José? *"
    p3 = "3) ¿Qué objeto mágico tenía José en el cuento? *"
    p4 = "4) Escribe el nombre de un animal que apareció en el océano de José: *"
    p5 = "5) ¿Te gustó el cuento? Escribe SÍ o NO y por qué (muy cortito): *"
    
    c1 = "NIVEL 1: Nombra al personaje principal y el escenario concreto."
    c2 = "NIVEL 2: Identifica objetos clave y personajes secundarios (animales)."
    c3 = "NIVEL 3: Expresa agrado o desagrado hacia la historia de forma muy básica."

elif grado == "2do":
    p1 = "1) ¿Quién es José y con qué estaba jugando? *"
    p2 = "2) Describe con tus propias palabras cómo era la tina océano: *"
    p3 = "3) ¿Qué problema o susto tuvo José durante su juego? *"
    p4 = "4) ¿Cómo logró José solucionar su problema? *"
    p5 = "5) ¿Qué nos enseña este cuento sobre la imaginación? *"
    
    c1 = "NIVEL 1: Describe al personaje y características literales del escenario mágico."
    c2 = "NIVEL 2: Identifica el problema central (nudo) y la resolución del mismo."
    c3 = "NIVEL 3: Reconoce el valor de la imaginación o el mensaje central del cuento."

elif grado == "3ro":
    p1 = "1) Explica qué hacía José para que su tina se convirtiera en un océano: *"
    p2 = "2) Nombra dos características importantes del mundo que imaginó José: *"
    p3 = "3) ¿Por qué crees que José sintió miedo en una parte de la historia? *"
    p4 = "4) ¿De qué manera la imaginación ayudó al protagonista a vencer el aburrimiento? *"
    p5 = "5) Si tú tuvieras una tina mágica, ¿a qué lugar viajarías y qué harías allí? *"
    
    c1 = "NIVEL 1: Explica el proceso imaginativo del personaje y detalles del mundo creado."
    c2 = "NIVEL 2: Infiere emociones (miedo) y comprende la función del juego en la trama."
    c3 = "NIVEL 3: Reflexiona sobre la lectura proyectando su propia experiencia imaginativa."

minimos = obtener_minimos(grado if grado != "" else "1ro")

if grado != "":
    st.info(f"✨ ¡Perfecto! Eres de {grado} de Primaria. Aquí están tus preguntas:")
    
    st.subheader("🔎 NIVEL 1: RECORDANDO EL CUENTO")
    q1 = st.text_input(p1, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga_cat1 and not tiempo_agotado and not q1.strip(): st.markdown(css_error, unsafe_allow_html=True)

    q2 = st.text_input(p2, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga_cat1 and not tiempo_agotado and not q2.strip(): st.markdown(css_error, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("🕵️‍♂️ NIVEL 2: PISTAS SECRETAS")
    q3 = st.text_area(p3, disabled=bloquear_preguntas, height=100)
    palabras_q3 = len(q3.split())
    if st.session_state.intento_descarga_cat1 and not tiempo_agotado:
        if not q3.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q3 < minimos["q3"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Debes escribir al menos {minimos['q3']} palabras (tienes {palabras_q3}). ¡Tú puedes!</div>", unsafe_allow_html=True)

    q4 = st.text_area(p4, disabled=bloquear_preguntas, height=100)
    palabras_q4 = len(q4.split())
    if st.session_state.intento_descarga_cat1 and not tiempo_agotado:
        if not q4.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q4 < minimos["q4"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Debes escribir al menos {minimos['q4']} palabras (tienes {palabras_q4}).</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("💡 NIVEL 3: TU OPINIÓN IMPORTA")
    q5 = st.text_area(p5, disabled=bloquear_preguntas, height=100)
    palabras_q5 = len(q5.split())
    if st.session_state.intento_descarga_cat1 and not tiempo_agotado:
        if not q5.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q5 < minimos["q5"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Debes escribir al menos {minimos['q5']} palabras (tienes {palabras_q5}).</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.error("🚨 **ÚLTIMO PASO:** Cuando presiones el botón de abajo, habrás terminado.")

    # --- LÓGICA DE VALIDACIÓN ---
    faltantes_personales = not institucion.strip() or not nombre.strip() or nivel == "" or grado == "" or seccion == "" or not area_curso.strip()
    preguntas_vacias = not q1.strip() or not q2.strip() or not q3.strip() or not q4.strip() or not q5.strip()
    faltan_palabras = palabras_q3 < minimos["q3"] or palabras_q4 < minimos["q4"] or palabras_q5 < minimos["q5"]
    hay_spam = contiene_spam(q1) or contiene_spam(q2) or contiene_spam(q3) or contiene_spam(q4) or contiene_spam(q5)

    hay_errores = False
    if faltantes_personales: hay_errores = True
    if not tiempo_agotado and (preguntas_vacias or faltan_palabras or hay_spam): hay_errores = True

    if not st.session_state.intento_descarga_cat1 or hay_errores:
        if st.button("📥 Generar y Descargar mi Tarea", type="primary", use_container_width=True):
            st.session_state.intento_descarga_cat1 = True
            st.rerun()
        
        if st.session_state.intento_descarga_cat1 and hay_errores:
            if faltantes_personales: st.error("⚠️ ¡Huy! Te olvidaste de poner tus datos. Revisa arriba.")
            elif not tiempo_agotado:
                if preguntas_vacias: st.warning("⚠️ Te falta responder algunas preguntas de tu cuento.")
                if faltan_palabras: st.warning("⚠️ Tus respuestas son muy cortitas. ¡Escribe un poquito más!")
                if hay_spam: st.error("⚠️ Estás escribiendo letras repetidas. Escribe palabras reales.")

    else:
        with st.spinner("Guardando tu tarea..."):
            duplicado = ya_participo(nombre, institucion, nombre_de_la_obra)
            
        if duplicado:
            st.error(f"🛑 **AVISO:** El alumno '{nombre.upper()}' del colegio '{institucion.upper()}' ya entregó esta tarea.")
        else:
            if tiempo_agotado:
                st.info("⏱️ El tiempo se acabó. Descargando lo que hiciste...")
            else:
                st.success("✨ ¡Súper! Haz clic en el botón rojo de abajo para descargar tu documento.")
                
            doc = Document()
            doc.add_heading('EVALUACIÓN DEL PLAN LECTOR (PRIMARIA)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'I.E.: {institucion.upper()}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'Obra: {nombre_de_la_obra} | Autor: Eduardo Florez Montero').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            fecha_str = fecha_input.strftime("%d/%m/%Y")
            texto_area = area_curso.strip()
            texto_prof = profesor.strip() if profesor.strip() else "No especificado"

            doc.add_paragraph(f'Estudiante: {nombre}\nNivel: {nivel} | Grado: {grado} | Sección: {seccion} | Fecha: {fecha_str}')
            doc.add_paragraph(f'Área/Curso: {texto_area} | Docente a cargo: {texto_prof}')
            
            def agregar_bloque_justificado(pregunta, respuesta):
                p_pregunta = doc.add_paragraph()
                p_pregunta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_pregunta.add_run(pregunta).bold = True
                p_respuesta = doc.add_paragraph()
                p_respuesta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                texto_respuesta = respuesta if respuesta.strip() else "[No respondió]"
                p_respuesta.add_run(f'Respuesta: {texto_respuesta}\n')

            doc.add_heading('NIVEL 1 (Literal)', level=2)
            agregar_bloque_justificado(p1, q1)
            agregar_bloque_justificado(p2, q2)
            
            doc.add_heading('NIVEL 2 (Inferencia)', level=2)
            agregar_bloque_justificado(p3, q3)
            agregar_bloque_justificado(p4, q4)
            
            doc.add_heading('NIVEL 3 (Crítico)', level=2)
            agregar_bloque_justificado(p5, q5)
            
            doc.add_page_break()
            doc.add_heading(f'Lista de Cotejo - Docente ({grado} Primaria)', level=2)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = 'CRITERIOS DE EVALUACIÓN', 'INICIO', 'PROCESO', 'LOGRADO', 'DESTACADO'
            
            criterios = [c1, c2, c3] 
            
            for crit in criterios:
                row_cells = table.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text, row_cells[4].text = crit, "[   ]", "[   ]", "[   ]", "[   ]"
                
            doc.add_paragraph('\nObservaciones / Feedback al estudiante:\n________________________________________________')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 Confirmar y Descargar mi Tarea", 
                data=bio.getvalue(), 
                file_name=f"TinaOceano_{grado}_{seccion}_{nombre.replace(' ', '_')}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                on_click=registrar_y_bloquear,
                args=(nombre, institucion, nivel, grado, seccion, area_curso, profesor, nombre_de_la_obra),
                type="primary",
                use_container_width=True
            )
else:
    st.info("👆 Por favor, diles a tus papis que te ayuden a seleccionar tu **Nivel** y **Grado** en la parte de arriba.")
