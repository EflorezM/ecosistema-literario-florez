import streamlit as st
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import streamlit.components.v1 as components
import json
import gspread

st.set_page_config(page_title="Chepeconde - Plan Lector", page_icon="📖", layout="centered")

# --- CONEXIÓN A GOOGLE SHEETS Y CANDADO DE 3 CERROJOS ---
@st.cache_resource
def get_google_sheet():
    creds_json = st.secrets["google_credentials"]
    creds_dict = json.loads(creds_json)
    if isinstance(creds_dict, list):
        creds_dict = creds_dict[0]
    client = gspread.service_account_from_dict(creds_dict)
    return client.open("Registro_Plan_Lector").sheet1

def ya_participo(nombre, inst, obra):
    try:
        sheet = get_google_sheet()
        data = sheet.get_all_values()
        if not data: return False
        for row in data[1:]:
            if len(row) >= 9:
                if (row[1].strip().lower() == nombre.strip().lower() and 
                    row[2].strip().lower() == inst.strip().lower() and 
                    row[8].strip().lower() == obra.strip().lower()):
                    return True
        return False
    except Exception as e:
        return False 

def registrar_y_bloquear(nombre, inst, nivel, grado, seccion, area, prof, obra):
    st.session_state.ficha_completada = True
    try:
        sheet = get_google_sheet()
        fecha_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        sheet.append_row([fecha_str, nombre.upper(), inst.upper(), nivel, grado, seccion, area.upper(), prof, obra.upper(), "COMPLETADO"])
    except Exception as e:
        pass

# --- FUNCIONES DE ESTADO Y VALIDACIÓN ---
if 'ficha_completada' not in st.session_state: st.session_state.ficha_completada = False
if 'intento_descarga' not in st.session_state: st.session_state.intento_descarga = False

if st.session_state.ficha_completada:
    st.error("🔒 **EVALUACIÓN ENTREGADA Y BLOQUEADA**")
    st.success("Has generado tu evidencia con éxito y tu participación ha sido registrada en el sistema central.")
    st.info("Por favor, envía el archivo Word que se descargó a tu profesor.")
    st.stop()

def contiene_spam(texto):
    for palabra in texto.split():
        if len(palabra) > 20: return True
    return False

def obtener_minimos(grado):
    if grado == "1ro": return {"q3": 10, "q4": 10, "q5": 10}
    elif grado == "2do": return {"q3": 12, "q4": 12, "q5": 12}
    elif grado == "3ro": return {"q3": 15, "q4": 15, "q5": 15}
    return {"q3": 10, "q4": 10, "q5": 10}

# --- 1. LA BÓVEDA DE SEGURIDAD ---
nombre_de_la_obra = "El Pueblo de Chepeconde"

if 'autenticado_chepe' not in st.session_state: st.session_state.autenticado_chepe = False

if not st.session_state.autenticado_chepe:
    st.title("📖 Ecosistema Literario Transmedia")
    st.subheader(f"Obra: {nombre_de_la_obra}")
    st.write("Autor: Eduardo Florez Montero")
    st.markdown("---")
    st.info("🔒 **Acceso Exclusivo:** Escribe la Clave de Aula que te indicó tu profesor.")
    
    clave = st.text_input("🔑 Clave de Acceso *:", type="password", autocomplete="new-password")
    
    if st.button("Desbloquear Evaluación", use_container_width=True):
        claves_validas = ["CHEPECONDE-2026", "AGUSTIN-3A", "AGUSTIN-3B", "FATIMA-2C", "BELEM-UNICA"]
        if clave in claves_validas: 
            st.session_state.autenticado_chepe = True
            st.rerun()
        else:
            st.error("❌ Clave incorrecta. Consulta con tu profesor.")
    st.stop()

# --- 2. SISTEMA DE TIEMPO INTELIGENTE ---
if 'ficha_iniciada_chepe' not in st.session_state: st.session_state.ficha_iniciada_chepe = False
if 'extra_time_used' not in st.session_state: st.session_state.extra_time_used = False

if not st.session_state.ficha_iniciada_chepe:
    st.success("✅ ¡Clave aceptada! Bienvenido al ecosistema.")
    st.warning("⚠️ **ATENCIÓN: TIENES UN SOLO INTENTO**\nEl sistema registrará tu nombre y colegio. No podrás volver a dar la evaluación de este libro.")
    
    try:
        st.image("portada_chepeconde.jpg", width=250)
    except:
        st.caption("(Aquí aparecerá la portada de tu libro)")
        
    st.info("👋 Tienes 20 minutos para resolver tu evaluación.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO, INICIAR EVALUACIÓN", use_container_width=True, type="primary"):
            st.session_state.ficha_iniciada_chepe = True
            st.session_state.inicio_tiempo_chepe = time.time()
            st.session_state.minutos_asignados_chepe = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_chepe
segundos_restantes = (st.session_state.minutos_asignados_chepe * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_preguntas = tiempo_agotado 

with st.sidebar:
    st.markdown("### ⏱️ Tiempo Restante")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <div id="aviso_30s" style="color: #e74c3c; font-weight: bold; font-size: 14px; text-align: center; margin-top: 10px; display: none;">
            ⚠️ ¡Aviso! Te quedan 30 segundos.
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var aviso = document.getElementById('aviso_30s');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                    aviso.style.display = "block";
                    aviso.innerHTML = "⚠️ Tiempo agotado.";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; 
                    if (tiempo <= 60) display.style.color = "#e74c3c"; 
                    if (tiempo <= 30) aviso.style.display = "block";
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=130)
        
        if not st.session_state.extra_time_used:
            if st.button("➕ Dar 4 min extra", use_container_width=True):
                st.session_state.minutos_asignados_chepe += 4
                st.session_state.extra_time_used = True
                st.rerun()
        else:
            st.caption("✔️ Tiempo extra utilizado.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Aún puedes descargar tu avance.")

# --- 3. CONTENIDO LITERARIO Y VALIDACIONES ---
st.markdown(f"### 📖 {nombre_de_la_obra.upper()}")
st.markdown("**Autor:** Eduardo Florez Montero | **Categoría:** Pre-Adolescentes (1ro a 3ro Sec.)")
st.markdown("---")

css_error = "<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Este campo es obligatorio</div>"

institucion = st.text_input("Institución Educativa (Nombre completo) *:", disabled=False)
if st.session_state.intento_descarga and not institucion.strip(): st.markdown(css_error, unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1: 
    nombre = st.text_input("Estudiante (Apellidos y Nombres) *:", disabled=False)
    if st.session_state.intento_descarga and not nombre.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col2: 
    nivel = st.selectbox("Nivel *:", ["", "Secundaria"], disabled=False)
    if st.session_state.intento_descarga and nivel == "": st.markdown(css_error, unsafe_allow_html=True)
with col3: 
    grado = st.selectbox("Grado *:", ["", "1ro", "2do", "3ro"], disabled=False)
    if st.session_state.intento_descarga and grado == "": st.markdown(css_error, unsafe_allow_html=True)

col4, col5, col6 = st.columns(3)
with col4: 
    seccion = st.selectbox("Sección *:", ["", "A", "B", "C", "D", "E", "F", "G", "H", "Única"], disabled=False)
    if st.session_state.intento_descarga and seccion == "": st.markdown(css_error, unsafe_allow_html=True)
with col5: 
    area_curso = st.text_input("Área o Curso (Ej: Comunicación) *:", disabled=False)
    if st.session_state.intento_descarga and not area_curso.strip(): st.markdown(css_error, unsafe_allow_html=True)
with col6: 
    fecha_input = st.date_input("Fecha *:", disabled=False)

profesor = st.text_input("Nombre de tu Profesor(a) (Opcional):", disabled=False)

st.markdown("---")
st.subheader("🎬 Antes de empezar: Resumen Visual")
st.video("https://www.youtube.com/watch?v=123456789") 
st.markdown("---")

# --- LÓGICA DE DIFERENCIACIÓN COGNITIVA Y RÚBRICA (PIAGET) ---
if grado == "1ro":
    p1 = "1) ¿Quién es el personaje principal y qué hace en el primer capítulo? *"
    p2 = "2) Describe brevemente cómo es El Pueblo de Chepeconde: *"
    p3 = "3) ¿Qué problema tuvo el protagonista y cómo lo resolvió? *"
    p4 = "4) ¿Quién consideras que es el antagonista (el que se opone al protagonista)? *"
    p5 = "5) ¿Qué enseñanza sencilla te deja la historia de Chepeconde? *"
    
    c1 = "NIVEL 1: Identifica personajes principales y describe escenarios basándose en detalles literales de la obra."
    c2 = "NIVEL 2: Deduce el problema principal y reconoce las acciones del antagonista."
    c3 = "NIVEL 3: Extrae una enseñanza clara y la relaciona con situaciones cotidianas."

elif grado == "2do":
    p1 = "1) Identifica al personaje principal y describe su mayor motivación: *"
    p2 = "2) ¿Qué características del pueblo de Chepeconde influyen en la historia? *"
    p3 = "3) ¿Por qué crees que el protagonista tomó esa decisión difícil? ¿Qué hubieras hecho tú? *"
    p4 = "4) Identifica el conflicto principal. ¿Consideras que es un conflicto interno o externo? *"
    p5 = "5) Reflexión: ¿Cómo se relaciona la historia de Chepeconde con los problemas de nuestra sociedad actual? *"
    
    c1 = "NIVEL 1: Caracteriza al personaje principal identificando su motivación y el contexto del pueblo."
    c2 = "NIVEL 2: Analiza los motivos detrás de decisiones difíciles y clasifica correctamente el conflicto."
    c3 = "NIVEL 3: Argumenta una reflexión ética vinculando los sucesos de la obra con la sociedad actual."

elif grado == "3ro":
    p1 = "1) Analiza al personaje principal: ¿Cuáles son sus fortalezas y debilidades psicológicas? *"
    p2 = "2) ¿De qué manera el entorno de Chepeconde condiciona el comportamiento de sus habitantes? *"
    p3 = "3) Evalúa el dilema ético del protagonista: ¿Justificas sus acciones finales? Argumenta tu respuesta. *"
    p4 = "4) Desglosa el conflicto principal y explica cómo afecta a los personajes secundarios. *"
    p5 = "5) Valoración crítica: Propón un final alternativo que cambie el mensaje social de la obra. *"
    
    c1 = "NIVEL 1: Analiza la psicología del personaje principal y cómo el entorno condiciona el comportamiento social."
    c2 = "NIVEL 2: Evalúa dilemas éticos y desglosa el impacto del conflicto central en personajes secundarios."
    c3 = "NIVEL 3: Formula juicios críticos avanzados, proponiendo hipótesis o finales alternativos lógicos."

minimos = obtener_minimos(grado if grado != "" else "1ro")

if grado != "":
    st.info(f"🧠 Evaluación adaptada al desarrollo cognitivo de {grado} de Secundaria.")
    
    # MOSTRAR RÚBRICA AL ESTUDIANTE
    st.markdown("### 📋 Rúbrica de Evaluación: ¿Qué evaluará tu profesor?")
    st.success(f"**NIVEL 1 (Literal):** {c1}\n\n**NIVEL 2 (Inferencial):** {c2}\n\n**NIVEL 3 (Crítico):** {c3}")
    st.markdown("---")
    
    st.subheader("NIVEL 1: COMPRENSIÓN LITERAL")
    q1 = st.text_input(p1, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga and not tiempo_agotado and not q1.strip(): st.markdown(css_error, unsafe_allow_html=True)

    q2 = st.text_area(p2, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga and not tiempo_agotado and not q2.strip(): st.markdown(css_error, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("NIVEL 2: ANÁLISIS E INFERENCIA")
    q3 = st.text_area(p3, disabled=bloquear_preguntas)
    palabras_q3 = len(q3.split())
    if st.session_state.intento_descarga and not tiempo_agotado:
        if not q3.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q3 < minimos["q3"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Mínimo {minimos['q3']} palabras (tienes {palabras_q3}).</div>", unsafe_allow_html=True)

    q4 = st.text_input(p4, disabled=bloquear_preguntas)
    if st.session_state.intento_descarga and not tiempo_agotado and not q4.strip(): st.markdown(css_error, unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("NIVEL 3: VALORACIÓN CRÍTICA")
    q5 = st.text_area(p5, disabled=bloquear_preguntas)
    palabras_q5 = len(q5.split())
    if st.session_state.intento_descarga and not tiempo_agotado:
        if not q5.strip(): st.markdown(css_error, unsafe_allow_html=True)
        elif palabras_q5 < minimos["q5"]: st.markdown(f"<div style='color:#e74c3c; font-size:13px; margin-top:-10px; margin-bottom:10px;'>⚠️ Mínimo {minimos['q5']} palabras (tienes {palabras_q5}).</div>", unsafe_allow_html=True)

    st.markdown("---")
    st.error("🚨 **ÚLTIMO PASO:** Al generar el Word, tu nombre quedará registrado en el sistema central.")

    # --- LÓGICA DE VALIDACIÓN ---
    faltantes_personales = not institucion.strip() or not nombre.strip() or nivel == "" or grado == "" or seccion == "" or not area_curso.strip()
    preguntas_vacias = not q1.strip() or not q2.strip() or not q3.strip() or not q4.strip() or not q5.strip()
    faltan_palabras = palabras_q3 < minimos["q3"] or palabras_q5 < minimos["q5"]
    hay_spam = contiene_spam(q1) or contiene_spam(q2) or contiene_spam(q3) or contiene_spam(q4) or contiene_spam(q5)

    hay_errores = False
    if faltantes_personales: hay_errores = True
    if not tiempo_agotado and (preguntas_vacias or faltan_palabras or hay_spam): hay_errores = True

    if not st.session_state.intento_descarga or hay_errores:
        if st.button("📥 Generar y Descargar mi Evidencia", type="primary", use_container_width=True):
            st.session_state.intento_descarga = True
            st.rerun()
        
        if st.session_state.intento_descarga and hay_errores:
            if faltantes_personales: st.error("⚠️ Faltan datos obligatorios. Revisa las alertas en rojo arriba.")
            elif not tiempo_agotado:
                if preguntas_vacias: st.warning("⚠️ Te falta responder algunas preguntas.")
                if faltan_palabras: st.warning("⚠️ Algunas respuestas no cumplen con la extensión mínima.")
                if hay_spam: st.error("⚠️ Escribe con normalidad, sin caracteres repetidos.")

    else:
        with st.spinner("Validando usuario en la Base de Datos..."):
            duplicado = ya_participo(nombre, institucion, nombre_de_la_obra)
            
        if duplicado:
            st.error(f"🛑 **BLOQUEO:** El estudiante '{nombre.upper()}' de la institución '{institucion.upper()}' ya rindió la evaluación de '{nombre_de_la_obra}'.")
        else:
            if tiempo_agotado:
                st.info("⏱️ El tiempo se agotó. Descargando avance parcial...")
            else:
                st.success("✨ ¡Todo correcto! Haz clic abajo para descargar tu evaluación.")
                
            doc = Document()
            doc.add_heading('EVALUACIÓN DEL PLAN LECTOR', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'I.E.: {institucion.upper()}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'Obra: {nombre_de_la_obra} | Autor: Eduardo Florez Montero').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            fecha_str = fecha_input.strftime("%d/%m/%Y")
            texto_area = area_curso.strip()
            texto_prof = profesor.strip() if profesor.strip() else "No especificado"

            doc.add_paragraph(f'Estudiante: {nombre}\nNivel: {nivel} | Grado: {grado} | Sección: {seccion} | Fecha: {fecha_str}')
            doc.add_paragraph(f'Área/Curso: {texto_area} | Docente a cargo: {texto_prof}')
            
            doc.add_heading('NIVEL 1 (Literal)', level=2)
            doc.add_paragraph(p1).bold = True
            doc.add_paragraph(f'Respuesta: {q1 if q1.strip() else "[No respondió]"}\n')
            doc.add_paragraph(p2).bold = True
            doc.add_paragraph(f'Respuesta: {q2 if q2.strip() else "[No respondió]"}\n')
            
            doc.add_heading('NIVEL 2 (Inferencia)', level=2)
            doc.add_paragraph(p3).bold = True
            doc.add_paragraph(f'Respuesta: {q3 if q3.strip() else "[No respondió]"}\n')
            doc.add_paragraph(p4).bold = True
            doc.add_paragraph(f'Respuesta: {q4 if q4.strip() else "[No respondió]"}\n')
            
            doc.add_heading('NIVEL 3 (Crítico)', level=2)
            doc.add_paragraph(p5).bold = True
            doc.add_paragraph(f'Respuesta: {q5 if q5.strip() else "[No respondió]"}\n')
            
            doc.add_page_break()
            doc.add_heading(f'Lista de Cotejo ({grado} Secundaria)', level=2)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text, hdr_cells[4].text = 'CRITERIOS DE EVALUACIÓN', 'INICIO', 'PROCESO', 'LOGRADO', 'DESTACADO'
            
            criterios = [c1, c2, c3] # Aquí se insertan los criterios específicos del grado
            
            for crit in criterios:
                row_cells = table.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text, row_cells[4].text = crit, "[   ]", "[   ]", "[   ]", "[   ]"
                
            doc.add_paragraph('\nObservaciones / Feedback al estudiante:\n________________________________________________')
            
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 Confirmar Registro y Descargar Evidencia", 
                data=bio.getvalue(), 
                file_name=f"Chepeconde_{grado}_{seccion}_{nombre.replace(' ', '_')}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                on_click=registrar_y_bloquear,
                args=(nombre, institucion, nivel, grado, seccion, area_curso, profesor, nombre_de_la_obra),
                type="primary",
                use_container_width=True
            )
else:
    st.info("👆 Por favor, selecciona tu **Nivel** y **Grado** en la parte superior para cargar las preguntas de tu evaluación.")
